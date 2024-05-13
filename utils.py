from docxtpl import DocxTemplate
from typing import Dict, List
import os
from docx import Document
from copy import deepcopy
from docx.oxml import OxmlElement
from docx.shared import Pt

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


file_agreements_paths: List[str] = [
    "templates\\Согласие_на_обработку_Шаблон.docx",
    "templates\\Согласие_на_указание_Шаблон.docx",
]

file_report_path: str = "templates\\РЕФЕРАТ_Шаблон.docx"
file_code_path: str = "templates\\Идентифицирующие_ПрЭВМ_Шаблон.docx"
file_contract_path: str = "templates\\ДОГОВОР_с_авторами_Шаблон.docx"
file_form_path: str = "templates\\АНКЕТА_РИД_Шаблон.docx"
file_calculation_path: str = "templates\\Калькуляция_НМА_Шаблон.docx"
file_notification_path: str = "templates\\Уведомление_на_ОАП_Шаблон.docx"



def set_cell_background(cell, fill_color):
    """ Установить цвет фона ячейки.
    :param cell: объект ячейки docx.table._Cell
    :param fill_color: Цвет фона в шестнадцатеричном формате, например 'FFFF00' для желтого
    """
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color))
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_authors_to_table(
    doc: Document, authors_info: Dict[str, Dict[str, Dict[str, str]]]
) -> Document:
    # Получаем первую таблицу в документе
    table = doc.tables[0]

    # Рассчитываем проценты для каждого автора, округлённые до двух знаков после запятой
    num_authors = len(authors_info)
    percent_per_author = round(100 / num_authors, 2)

    # Чтобы сумма процентов была ровно 100, корректируем процент для последнего автора
    # Это компенсирует потери при округлении
    total_percent = percent_per_author * (num_authors - 1)
    last_author_percent = round(100 - total_percent, 2)

    # Добавляем информацию об авторах в таблицу
    for index, details in enumerate(authors_info.values()):
        # Формируем полное ФИО
        full_name = f"{details['subject_name']['surname']} {details['subject_name']['name']} {details['subject_name']['middle_name']}"
        # Добавляем строку
        row_cells = table.add_row().cells
        row_cells[0].text = full_name
        # Распределение процентов с корректировкой для последнего автора
        if index == num_authors - 1:  # Если это последний автор в списке
            row_cells[1].text = f"{last_author_percent}%"  # Корректируем процент
        else:
            row_cells[1].text = f"{percent_per_author}%"

    return doc


def format_address(address):
    """
    Форматирует словарь с информацией об адресе в строку в формате паспорта РФ.

    :param address: Словарь с ключами country, post_index, city, street, home_num, appartment_num
    :return: Строка с форматированным адресом
    """
    formatted_address = (
        f"{address['country']}, {address['post_index']}, г. {address['city']}, "
        f"ул. {address['street']}, д. {address['home_num']}, кв. {address['appartment_num']}"
    )
    return formatted_address


def format_author_signature(authors_info):
    """
    Формирует информацию о подписи авторов для документа.

    :param authors_info: Словарь с информацией об авторах.
    :return: Строка, содержащая информацию о подписях авторов, разделенная символом новой строки.
    """
    template = (
        "{surname} {name} {middle_name}\n_______________ / {initials} {surname}\n"
    )
    author_texts = []
    for details in authors_info.values():
        initials = f"{details['subject_name']['name'][0]}.{details['subject_name']['middle_name'][0]}."
        text = template.format(
            surname=details["subject_name"]["surname"],
            name=details["subject_name"]["name"],
            middle_name=details["subject_name"]["middle_name"],
            initials=initials,
        )
        author_texts.append(text)
    return "\n".join(author_texts)


def format_name(subject_name, full_name=True):
    """
    Форматирует информацию о личности в строку.

    :param subject_name: Словарь с ключами surname, name, middle_name, содержащий личные данные.
    :param full_name: Булевый параметр, который определяет формат вывода:
                      True - полное имя, False - фамилия и инициалы.
    :return: Строка с форматированным именем.
    """
    if full_name:
        # Форматирование полного имени
        return f"{subject_name['surname']} {subject_name['name']} {subject_name['middle_name']}"
    else:
        # Форматирование фамилии и инициалов
        initials = f"{subject_name['name'][0]}.{subject_name['middle_name'][0]}."
        return f"{subject_name['surname']} {initials}"


def format_authors_info(authors_info):
    """
    Форматирует информацию об авторах по заданному шаблону.

    :param authors_info: Словарь с информацией об авторах.
    :return: Строка, содержащая информацию обо всех авторах, разделенная символом новой строки.
    """
    template = (
        "{surname}, {name}, {middle_name}, личность удостоверена паспортом серии {series} № {number}, "
        "выданным {issued_by}, код подразделения {division_code}, дата выдачи {issue_date}, "
        "зарегистрированный по адресу {post_index}, г. {city}, дом {home_num}, квартира {appartment_num}, "
        "являющийся (являющаяся) сотрудником Университета, в дальнейшем именуемый «Соавтор»;"
    )
    author_texts = []
    for details in authors_info.values():
        text = template.format(
            surname=details["subject_name"]["surname"],
            name=details["subject_name"]["name"],
            middle_name=details["subject_name"]["middle_name"],
            series=details["passport_series"],
            number=details["passport_number"],
            issued_by=details["passport_issued_by"],
            division_code=details["passport_division_code"],
            issue_date=details["passport_issue_date"],
            post_index=details["subject_address"]["post_index"],
            city=details["subject_address"]["city"],
            home_num=details["subject_address"]["home_num"],
            appartment_num=details["subject_address"]["appartment_num"],
        )
        author_texts.append(text)
    return "\n".join(author_texts)


# Утилита для сохранения документа
def render_document(
    template_path: str, replacements: Dict[str, str], save_path: str, identifier: str
) -> None:
    doc = DocxTemplate(template_path)
    doc.render(replacements)

    file_name = os.path.basename(template_path.replace("Шаблон", identifier))
    full_file_path = os.path.join(save_path, file_name)
    doc.save(full_file_path)


# Соглашения
def render_agreements(
    replacements_autors: Dict[str, Dict[str, str]],
    replacements_programm: Dict[str, str],
    save_path: str,
) -> None:
    for file_agreement_path in file_agreements_paths:
        for replacements in replacements_autors.values():
            formatted_address = {
                "formatted_address": format_address(replacements["subject_address"])
            }
            formatted_name_short = {
                "subject_name_short": format_name(
                    replacements["subject_name"], full_name=False
                )
            }
            formatted_name_long = {
                "subject_name_long": format_name(replacements["subject_name"])
            }

            render_document(
                file_agreement_path,
                replacements
                | replacements_programm
                | formatted_address
                | formatted_name_short
                | formatted_name_long,
                save_path,
                formatted_name_short["subject_name_short"],
            )


# Реферат
def render_report(
    replacements_autors: Dict[str, Dict[str, str]],
    replacements_programm: Dict[str, str],
    save_path: str,
) -> None:
    author_names = ", ".join(
        format_name(replacements["subject_name"], full_name=True)
        for replacements in replacements_autors.values()
    )
    replacements = {"author_names_long": author_names} | replacements_programm
    render_document(
        file_report_path, replacements, save_path, replacements_programm["theme"]
    )


# Идентифицирующие материалы
def render_code(
    replacements_autors: Dict[str, Dict[str, str]],
    replacements_programm: Dict[str, str],
    save_path: str,
) -> None:

    author_names = "\n ".join(
        format_name(replacements["subject_name"], full_name=False)
        for replacements in replacements_autors.values()
    )

    replacements = {"author_names_short": author_names} | replacements_programm
    render_document(
        file_code_path, replacements, save_path, replacements_programm["theme"]
    )


def render_contract(
    replacements_autors: Dict[str, Dict[str, str]],
    replacements_programm: Dict[str, str],
    save_path: str,
):
    authors_info = {"authors_info": format_authors_info(replacements_autors)}
    author_signature = {
        "authors_signature": format_author_signature(replacements_autors)
    }

    replacements = authors_info | replacements_programm | author_signature

    doc = DocxTemplate(file_contract_path)
    doc.render(replacements)

    file_name = os.path.basename(
        file_contract_path.replace("Шаблон", replacements_programm["theme"])
    )
    full_file_path = os.path.join(save_path, file_name)

    add_authors_to_table(doc, replacements_autors)
    doc.save(full_file_path)


def render_form(
    replacements_autors: Dict[str, Dict[str, str]],
    replacements_programm: Dict[str, str],
    save_path: str,
):
    render_document(
        file_form_path, replacements_programm, save_path, replacements_programm["theme"]
    )


def render_calculation(
    replacements_autors: Dict[str, Dict[str, str]],
    replacements_programm: Dict[str, str],
    save_path: str,
):
    doc = Document(file_calculation_path)

    reasons_to_table_index = {
        "в силу закона": 0,
        "в силу договора": 1,
        "в силу свободного": 2,
    }

    new_tables = []  # Список для хранения новых таблиц

    # Добавляем новые таблицы в документ
    for author_data in replacements_autors.values():
        reason = author_data['reason']
        if reason in reasons_to_table_index:
            table_index = reasons_to_table_index[reason]
            original_table = doc.tables[table_index]
            new_table_xml = deepcopy(original_table._tbl)
            doc._body._element.append(new_table_xml)
            new_table = doc.tables[-1]
            new_tables.append(new_table)  # Сохраняем объект таблицы, а не XML

    # Удаляем оригинальные таблицы
    original_tables_to_remove = [doc.tables[i] for i in reasons_to_table_index.values()]
    for table in original_tables_to_remove:
        p_element = table._element.getparent()
        p_element.remove(table._element)

    # Заполняем новые таблицы данными авторов и добавляем пустые параграфы
    for author_data, new_table in zip(replacements_autors.values(), new_tables):
        full_name = f"{author_data['subject_name']['surname']} {author_data['subject_name']['name']} {author_data['subject_name']['middle_name']}"
        cell = new_table.rows[0].cells[1]
        para = cell.paragraphs[0]
        run = para.add_run(full_name)
        run.font.size = Pt(10)  # Устанавливаем размер шрифта 10 пунктов
        para.add_run()  # Завершаем текущий run, чтобы применить стиль

        # Добавляем пустой параграф после каждой таблицы
        p = doc.add_paragraph()
        new_table._element.addnext(p._p)




    file_name = os.path.basename(file_calculation_path.replace("Шаблон", replacements_programm["theme"]))
    full_file_path = os.path.join(save_path, file_name)

    # Сохранение измененного документа
    doc.save(full_file_path)


def render_notification(
    replacements_autors: Dict[str, Dict[str, str]],
    replacements_programm: Dict[str, str],
    save_path: str,
):
    doc = Document(file_notification_path)

    target_table = doc.tables[2]

    template_rows = [row for row in target_table.rows]

    reasons_to_row_index = {
        "в силу закона": 0,
        "в силу договора": 1,
        "в силу свободного": 2,
    }

    for author_data in replacements_autors.values():
        reason = author_data['reason']
        if reason in reasons_to_row_index:
            row_index = reasons_to_row_index[reason]
            template_row = template_rows[row_index]
            new_row = target_table.add_row()
            for idx, cell in enumerate(new_row.cells):
                set_cell_background(cell, 'D9D9D9')  # Пример установки серого фона
                for para in template_row.cells[idx].paragraphs:
                    new_para = cell.add_paragraph()
                    new_para.paragraph_format.alignment = para.paragraph_format.alignment
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        # Копируем стили шрифта
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.size = run.font.size
                        new_run.font.color.rgb = run.font.color.rgb
                        new_run.font.name = run.font.name

    file_name = os.path.basename(file_notification_path.replace("Шаблон", replacements_programm["theme"]))
    full_file_path = os.path.join(save_path, file_name)

    # Сохранение измененного документа
    doc.save(full_file_path)